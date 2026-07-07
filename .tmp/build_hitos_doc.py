from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = Path("Informe_ejecutivo_hitos_realizados.docx")
BLUE = "1F4E78"
DARK = "243746"
MUTED = "5B6573"
LIGHT = "EAF1F7"
PALE = "F6F8FA"
WHITE = "FFFFFF"


HITOS = [
    {
        "numero": 1,
        "titulo": "Base del sistema",
        "periodo": "Semanas 1 y 2",
        "objetivo": "Poner en marcha una plataforma ordenada, segura y preparada para crecer.",
        "logros": [
            "Se habilitó el acceso al sistema y la navegación principal.",
            "Se estableció una estructura común para desarrollar los siguientes módulos de forma consistente.",
            "Se incorporaron respuestas claras ante cargas, errores y acciones completadas.",
        ],
        "resultado": "El proyecto quedó operativo y listo para incorporar las funciones del negocio sin rehacer su base.",
    },
    {
        "numero": 2,
        "titulo": "Usuarios, roles y configuración",
        "periodo": "Semanas 3 y 4",
        "objetivo": "Definir quién puede ingresar al sistema y qué acciones puede realizar.",
        "logros": [
            "Se organizaron los perfiles de usuario y los roles de trabajo.",
            "Se establecieron accesos diferenciados según las responsabilidades de cada usuario.",
            "Se habilitó una configuración general y una base de trazabilidad de acciones relevantes.",
        ],
        "resultado": "La plataforma quedó preparada para operar con responsabilidades claras y accesos controlados.",
    },
    {
        "numero": 3,
        "titulo": "Catálogos y reglas maestras",
        "periodo": "Semanas 5 y 6",
        "objetivo": "Centralizar los valores y criterios que utiliza la operación diaria.",
        "logros": [
            "Se organizaron estados, categorías, tipos de actividad, tamaños de empresa y otras opciones de uso frecuente.",
            "Se dejaron listas las opciones que alimentan formularios y procesos posteriores.",
            "Se redujo la dependencia de valores fijos, facilitando ajustes futuros.",
        ],
        "resultado": "El sistema quedó alineado con reglas comunes y parámetros reutilizables en todos sus módulos.",
    },
    {
        "numero": 4,
        "titulo": "Gestión de prospectos",
        "periodo": "Semanas 7 y 8",
        "objetivo": "Gestionar de principio a fin a las organizaciones interesadas en asociarse.",
        "logros": [
            "Se habilitó el registro, búsqueda, consulta y actualización de prospectos.",
            "Se incorporaron la evaluación, la cotización y el seguimiento de cada cambio de estado.",
            "Se registró al captador responsable y se preparó el paso hacia la conversión en asociado.",
        ],
        "resultado": "La organización puede dar seguimiento ordenado a cada oportunidad desde su ingreso hasta su aprobación.",
    },
    {
        "numero": 5,
        "titulo": "Conversión y ficha del asociado",
        "periodo": "Semanas 9 y 10",
        "objetivo": "Convertir prospectos aprobados y administrar la información central de cada asociado.",
        "logros": [
            "Se habilitó la conversión controlada de un prospecto aprobado en asociado.",
            "Se creó una ficha central con información empresarial y datos de contacto.",
            "Se incorporó la gestión de personas vinculadas y contactos organizados por área.",
        ],
        "resultado": "El sistema cuenta con un registro único y trazable para administrar la relación con cada asociado.",
    },
    {
        "numero": 6,
        "titulo": "Membresías, pagos y cobranza",
        "periodo": "Semanas 11 y 12",
        "objetivo": "Controlar el ciclo económico y administrativo de la membresía.",
        "logros": [
            "Se habilitó la creación de membresías y la programación de sus cuotas.",
            "Se incorporó el registro de pagos y su relación con obligaciones pendientes.",
            "Se organizó el seguimiento de cobranza, incluyendo resultados y próximas acciones.",
        ],
        "resultado": "La organización puede conocer la situación de pago de cada asociado y dar seguimiento oportuno a la cobranza.",
    },
    {
        "numero": 7,
        "titulo": "Gestión documental",
        "periodo": "Semanas 13 y 14",
        "objetivo": "Centralizar y ordenar los documentos relacionados con la operación y los asociados.",
        "logros": [
            "Se habilitó la carga, clasificación, búsqueda y descarga de documentos.",
            "Se vinculó la documentación con cada asociado y con otros contextos de la operación.",
            "Se incorporaron controles para conservar versiones y facilitar una consulta segura.",
        ],
        "resultado": "La información documental quedó centralizada, accesible y vinculada con los procesos correspondientes.",
    },
    {
        "numero": 8,
        "titulo": "Reportes, exportaciones y automatizaciones",
        "periodo": "Semanas 15 y 16",
        "objetivo": "Convertir la información registrada en herramientas de seguimiento y decisión.",
        "logros": [
            "Se habilitaron reportes de prospectos, asociados, membresías, pagos, cobranza y documentos.",
            "Se incorporaron indicadores y vistas resumen para una lectura rápida de la operación.",
            "Se habilitó la exportación de información y se dejó una base para futuras alertas y tareas automáticas.",
        ],
        "resultado": "La organización puede supervisar su operación, analizar resultados y compartir información consolidada.",
    },
]


def set_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_borders(table, color="D5DDE5", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def border_bottom(paragraph, color=BLUE, size="14", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, end])


def configure_bullets(doc):
    numbering = doc.part.numbering_part.element
    a_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    n_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    aid, nid = max(a_ids, default=0) + 1, max(n_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(aid))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    for tag, value in (("start", "1"), ("numFmt", "bullet"), ("lvlText", "•"), ("lvlJc", "left")):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:val"), value)
        lvl.append(el)
    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "650")
    ind.set(qn("w:hanging"), "300")
    ppr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "100")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    lvl.append(ppr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(nid))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(aid))
    num.append(ref)
    numbering.append(num)
    return nid


def add_bullet(doc, num_id, text):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    r = p.add_run(text)
    set_font(r)


def add_result_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "5")
        border.set(qn("w:color"), "B8CADD")
        borders.append(border)
    p_pr.append(borders)
    r = p.add_run("Resultado alcanzado: ")
    set_font(r, color=BLUE, bold=True)
    r = p.add_run(text)
    set_font(r, color=DARK)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in (
    ("Title", 28, DARK, 0, 8),
    ("Subtitle", 14, MUTED, 0, 16),
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
):
    style = doc.styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = name != "Subtitle"
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.text = "SISTEMA DE ASOCIADOS  |  INFORME EJECUTIVO"
for run in header.runs:
    set_font(run, size=8.5, color=MUTED, bold=True)
border_bottom(header, color="C8D2DC", size="4", space="4")
page_number(section.footer.paragraphs[0])
bullet_id = configure_bullets(doc)

# Portada
space = doc.add_paragraph()
space.paragraph_format.space_after = Pt(58)
kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
kicker.paragraph_format.space_after = Pt(14)
r = kicker.add_run("INFORME EJECUTIVO")
set_font(r, size=10, color=BLUE, bold=True)
title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Hitos realizados\ndel Sistema de Asociados")
subtitle = doc.add_paragraph(style="Subtitle")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("Ocho hitos ejecutados en periodos de dos semanas")
rule = doc.add_paragraph()
rule.paragraph_format.space_before = Pt(14)
rule.paragraph_format.space_after = Pt(28)
border_bottom(rule, color=BLUE, size="18", space="1")
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Duración total referencial: 16 semanas")
set_font(r, size=11, color=DARK, bold=True)
date = doc.add_paragraph()
date.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = date.add_run("30 de junio de 2026")
set_font(r, size=10.5, color=MUTED)
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
note.paragraph_format.space_before = Pt(38)
note.paragraph_format.left_indent = Inches(0.55)
note.paragraph_format.right_indent = Inches(0.55)
r = note.add_run(
    "El informe presenta exclusivamente los ocho hitos principales definidos para el proyecto. "
    "No incluye actividades ni hitos de subsanación."
)
set_font(r, size=10.5, color=MUTED, italic=True)

doc.add_page_break()
doc.add_heading("Resumen ejecutivo", level=1)
doc.add_paragraph(
    "El Sistema de Asociados fue desarrollado mediante ocho hitos consecutivos. Cada hito tuvo una duración de dos "
    "semanas y entregó una capacidad concreta, desde la puesta en marcha de la plataforma hasta la generación de "
    "reportes para el seguimiento y la toma de decisiones."
)
doc.add_paragraph(
    "El resultado es una solución integrada que permite organizar el acceso de los usuarios, gestionar prospectos y "
    "asociados, controlar membresías y pagos, centralizar documentos y consultar información consolidada."
)

table = doc.add_table(rows=1, cols=3)
for idx, text in enumerate(("Hito", "Periodo", "Resultado principal")):
    cell = table.rows[0].cells[idx]
    shade(cell, BLUE)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, size=9.5, color=WHITE, bold=True)
header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
header_flag = OxmlElement("w:tblHeader")
header_flag.set(qn("w:val"), "true")
header_tr_pr.append(header_flag)
for idx, hito in enumerate(HITOS):
    cells = table.add_row().cells
    values = (str(hito["numero"]), hito["periodo"], hito["titulo"])
    for col, text in enumerate(values):
        if idx % 2:
            shade(cells[col], PALE)
        p = cells[col].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col < 2 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        set_font(r, size=9.5, color=DARK, bold=(col == 0))
table_geometry(table, [900, 2100, 6360])
table_borders(table)

doc.add_heading("Resultado global", level=2)
for text in (
    "Una ruta completa desde la captación de un prospecto hasta la administración del asociado.",
    "Mayor control sobre membresías, pagos pendientes y acciones de cobranza.",
    "Documentación centralizada y relacionada con los procesos de la organización.",
    "Información resumida y exportable para supervisión y toma de decisiones.",
):
    add_bullet(doc, bullet_id, text)

# Dos hitos por página para conservar un formato ejecutivo.
for idx, hito in enumerate(HITOS):
    if idx % 2 == 0:
        doc.add_page_break()
    label = doc.add_paragraph()
    label.paragraph_format.space_after = Pt(2)
    r = label.add_run(f"HITO {hito['numero']}  |  {hito['periodo'].upper()}  |  REALIZADO")
    set_font(r, size=9.5, color=BLUE, bold=True)
    heading = doc.add_paragraph(style="Heading 1")
    heading.paragraph_format.space_before = Pt(0)
    heading.add_run(hito["titulo"])
    p = doc.add_paragraph()
    r = p.add_run("Propósito. ")
    set_font(r, bold=True)
    r = p.add_run(hito["objetivo"])
    set_font(r)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Principales avances")
    set_font(r, size=11, color=DARK, bold=True)
    for logro in hito["logros"]:
        add_bullet(doc, bullet_id, logro)
    add_result_box(doc, hito["resultado"])
    if idx % 2 == 0:
        separator = doc.add_paragraph()
        separator.paragraph_format.space_before = Pt(9)
        separator.paragraph_format.space_after = Pt(4)
        border_bottom(separator, color="D5DDE5", size="6", space="1")

doc.add_page_break()
doc.add_heading("Conclusión ejecutiva", level=1)
doc.add_paragraph(
    "Los ocho hitos conforman una secuencia completa de implementación. Las primeras seis semanas permitieron establecer "
    "la base de operación y sus reglas; las siguientes diez semanas incorporaron los procesos centrales de prospectos, "
    "asociados, finanzas, documentos y reportes."
)
doc.add_paragraph(
    "Al cierre de las 16 semanas referenciales, la organización cuenta con una plataforma que integra información, "
    "reduce tareas dispersas y mejora la visibilidad sobre la relación con prospectos y asociados."
)
add_result_box(
    doc,
    "Se completó el producto base previsto en los ocho hitos principales, con capacidades para operar, controlar y consultar la gestión de asociados.",
)
source = doc.add_paragraph()
source.paragraph_format.space_before = Pt(16)
r = source.add_run("Fuente de alcance: documentos de hitos ubicados en .agent/hitos/. ")
set_font(r, size=9, color=MUTED, italic=True)
r = source.add_run("Se excluyó la carpeta .agent/hitos/subsanacion/.")
set_font(r, size=9, color=MUTED, italic=True)

doc.core_properties.title = "Informe ejecutivo de hitos realizados"
doc.core_properties.subject = "Hitos 1 al 8 del Sistema de Asociados"
doc.core_properties.author = "Equipo del Sistema de Asociados"
doc.core_properties.keywords = "hitos, informe ejecutivo, sistema de asociados"
doc.save(OUT)
print(OUT.resolve())
