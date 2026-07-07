from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = Path("Informe_ejecutivo_hitos_realizados_APA7.docx")
FONT = "Times New Roman"
BLACK = "000000"

HITOS = [
    (1, "Base del sistema", "Semanas 1 y 2",
     "Poner en marcha una plataforma ordenada, segura y preparada para crecer.",
     ["Se habilitó el acceso al sistema y la navegación principal.",
      "Se estableció una estructura común para desarrollar los módulos posteriores de forma consistente.",
      "Se incorporaron respuestas claras ante cargas, errores y acciones completadas."],
     "El proyecto quedó operativo y listo para incorporar las funciones del negocio sin rehacer su base."),
    (2, "Usuarios, roles y configuración", "Semanas 3 y 4",
     "Definir quién puede ingresar al sistema y qué acciones puede realizar.",
     ["Se organizaron los perfiles de usuario y los roles de trabajo.",
      "Se establecieron accesos diferenciados según las responsabilidades de cada usuario.",
      "Se habilitó una configuración general y una base de trazabilidad para acciones relevantes."],
     "La plataforma quedó preparada para operar con responsabilidades claras y accesos controlados."),
    (3, "Catálogos y reglas maestras", "Semanas 5 y 6",
     "Centralizar los valores y criterios que utiliza la operación diaria.",
     ["Se organizaron estados, categorías, tipos de actividad, tamaños de empresa y otras opciones de uso frecuente.",
      "Se dejaron listas las opciones que alimentan formularios y procesos posteriores.",
      "Se redujo la dependencia de valores fijos, facilitando ajustes futuros."],
     "El sistema quedó alineado con reglas comunes y parámetros reutilizables en todos sus módulos."),
    (4, "Gestión de prospectos", "Semanas 7 y 8",
     "Gestionar de principio a fin a las organizaciones interesadas en asociarse.",
     ["Se habilitó el registro, búsqueda, consulta y actualización de prospectos.",
      "Se incorporaron la evaluación, la cotización y el seguimiento de cada cambio de estado.",
      "Se registró al captador responsable y se preparó el paso hacia la conversión en asociado."],
     "La organización puede dar seguimiento ordenado a cada oportunidad desde su ingreso hasta su aprobación."),
    (5, "Conversión y ficha del asociado", "Semanas 9 y 10",
     "Convertir prospectos aprobados y administrar la información central de cada asociado.",
     ["Se habilitó la conversión controlada de un prospecto aprobado en asociado.",
      "Se creó una ficha central con información empresarial y datos de contacto.",
      "Se incorporó la gestión de personas vinculadas y contactos organizados por área."],
     "El sistema cuenta con un registro único y trazable para administrar la relación con cada asociado."),
    (6, "Membresías, pagos y cobranza", "Semanas 11 y 12",
     "Controlar el ciclo económico y administrativo de la membresía.",
     ["Se habilitó la creación de membresías y la programación de sus cuotas.",
      "Se incorporó el registro de pagos y su relación con obligaciones pendientes.",
      "Se organizó el seguimiento de cobranza, incluyendo resultados y próximas acciones."],
     "La organización puede conocer la situación de pago de cada asociado y dar seguimiento oportuno a la cobranza."),
    (7, "Gestión documental", "Semanas 13 y 14",
     "Centralizar y ordenar los documentos relacionados con la operación y los asociados.",
     ["Se habilitó la carga, clasificación, búsqueda y descarga de documentos.",
      "Se vinculó la documentación con cada asociado y con otros contextos de la operación.",
      "Se incorporaron controles para conservar versiones y facilitar una consulta segura."],
     "La información documental quedó centralizada, accesible y vinculada con los procesos correspondientes."),
    (8, "Reportes, exportaciones y automatizaciones", "Semanas 15 y 16",
     "Convertir la información registrada en herramientas de seguimiento y decisión.",
     ["Se habilitaron reportes de prospectos, asociados, membresías, pagos, cobranza y documentos.",
      "Se incorporaron indicadores y vistas resumen para una lectura rápida de la operación.",
      "Se habilitó la exportación de información y se dejó una base para futuras alertas y tareas automáticas."],
     "La organización puede supervisar su operación, analizar resultados y compartir información consolidada."),
]


def font_run(run, bold=None, italic=None, size=12):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), FONT)
    rpr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(BLACK)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_apa_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8" if edge in ("top", "bottom") else "4")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    for edge in ("left", "right", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def page_field(run):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, display, end])


def configure_bullets(doc):
    numbering = doc.part.numbering_part.element
    aids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    nids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    aid, nid = max(aids, default=0) + 1, max(nids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(aid))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    for tag, val in (("start", "1"), ("numFmt", "bullet"), ("lvlText", "•"), ("lvlJc", "left")):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:val"), val)
        lvl.append(el)
    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "480")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    lvl.append(ppr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(nid))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(aid))
    num.append(ref)
    numbering.append(num)
    return nid


def add_bullet(doc, nid, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    ppr = p._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(nid))
    numpr.extend([ilvl, numid])
    ppr.append(numpr)
    font_run(p.add_run(text))


def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5) if indent else Inches(0)
    font_run(p.add_run(text))
    return p


def add_level1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_run(p.add_run(text), bold=True)
    return p


def add_level2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    font_run(p.add_run(text), bold=True)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.5)
section.footer_distance = Inches(0.5)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(12)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 2

for name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
    style = doc.styles[name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor.from_string(BLACK)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 2
    style.paragraph_format.keep_with_next = True

# APA 7 professional-paper running head and page number.
header = section.header.paragraphs[0]
header.paragraph_format.space_before = Pt(0)
header.paragraph_format.space_after = Pt(0)
header.paragraph_format.line_spacing = 1
header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
font_run(header.add_run("HITOS REALIZADOS DEL SISTEMA DE ASOCIADOS"), size=12)
font_run(header.add_run("\t"), size=12)
page_run = header.add_run()
font_run(page_run, size=12)
page_field(page_run)

bullet_id = configure_bullets(doc)

# Portada APA 7.
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(60)
title.paragraph_format.space_after = Pt(0)
title.paragraph_format.line_spacing = 2
font_run(title.add_run("Hitos realizados del Sistema de Asociados"), bold=True)
for text in (
    "Informe ejecutivo",
    "Equipo del Sistema de Asociados",
    "Proyecto Sistema de Asociados",
    "30 de junio de 2026",
):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    font_run(p.add_run(text))

doc.add_page_break()
add_level1(doc, "Resumen")
add_body(doc,
    "Este informe presenta los ocho hitos principales realizados para la implementación del Sistema de Asociados. "
    "Cada hito tuvo una duración de dos semanas, para un total referencial de 16 semanas. La secuencia comenzó con "
    "la preparación de la plataforma, los accesos y las reglas comunes; continuó con la gestión de prospectos, "
    "asociados, membresías, pagos y documentos; y concluyó con reportes, exportaciones y una base para futuras "
    "automatizaciones. El resultado es una solución integrada que mejora el orden de la información, el seguimiento "
    "de la operación y la toma de decisiones. El alcance excluye expresamente los hitos de subsanación.", indent=False)
keywords = doc.add_paragraph()
keywords.paragraph_format.first_line_indent = Inches(0.5)
font_run(keywords.add_run("Palabras clave: "), italic=True)
font_run(keywords.add_run("asociados, hitos, implementación, gestión, transformación digital"))

doc.add_page_break()
add_level1(doc, "Hitos realizados del Sistema de Asociados")
add_body(doc,
    "La implementación se organizó en ocho entregas consecutivas de dos semanas. Este enfoque permitió incorporar "
    "capacidades de manera progresiva y mantener una relación clara entre cada avance y el resultado esperado para "
    "la organización (Equipo del Sistema de Asociados, s. f.-a, s. f.-b, s. f.-c, s. f.-d, s. f.-e, s. f.-f, s. f.-g, s. f.-h).")

# Tabla APA.
label = doc.add_paragraph()
label.paragraph_format.first_line_indent = Inches(0)
font_run(label.add_run("Tabla 1"), bold=True)
caption = doc.add_paragraph()
caption.paragraph_format.first_line_indent = Inches(0)
font_run(caption.add_run("Cronograma referencial de los hitos realizados"), italic=True)
table = doc.add_table(rows=1, cols=3)
for idx, text in enumerate(("Hito", "Periodo", "Resultado principal")):
    p = table.rows[0].cells[idx].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx < 2 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = Pt(0)
    font_run(p.add_run(text), bold=True, size=10)
header_pr = table.rows[0]._tr.get_or_add_trPr()
header_flag = OxmlElement("w:tblHeader")
header_flag.set(qn("w:val"), "true")
header_pr.append(header_flag)
for number, title_text, period, *_ in HITOS:
    cells = table.add_row().cells
    for idx, text in enumerate((str(number), period, title_text)):
        p = cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx < 2 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_after = Pt(0)
        font_run(p.add_run(text), size=10)
set_table_geometry(table, [900, 2100, 6360])
set_apa_table_borders(table)
note = doc.add_paragraph()
note.paragraph_format.first_line_indent = Inches(0)
font_run(note.add_run("Nota. "), italic=True)
font_run(note.add_run("Cada hito tuvo una duración de dos semanas. El cronograma no incluye hitos de subsanación."))

add_level2(doc, "Resultado global")
add_body(doc,
    "En conjunto, los hitos habilitaron una ruta completa desde el ingreso de un prospecto hasta la administración "
    "del asociado, el control de sus compromisos, la organización documental y la consulta de indicadores. La "
    "duración total referencial del trabajo fue de 16 semanas.")

suffixes = "abcdefgh"
for idx, (number, title_text, period, objective, achievements, result) in enumerate(HITOS):
    if idx % 2 == 0:
        doc.add_page_break()
    add_level1(doc, f"Hito {number}: {title_text}")
    meta = doc.add_paragraph()
    meta.paragraph_format.first_line_indent = Inches(0)
    font_run(meta.add_run(f"Duración: {period}. Estado: realizado."), italic=True)
    add_body(doc, f"El propósito de este hito fue {objective[0].lower() + objective[1:]}")
    add_level2(doc, "Principales avances")
    for achievement in achievements:
        add_bullet(doc, bullet_id, achievement)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    font_run(p.add_run("Resultado alcanzado. "), bold=True)
    font_run(p.add_run(f"{result} (Equipo del Sistema de Asociados, s. f.-{suffixes[idx]})."))

doc.add_page_break()
add_level1(doc, "Conclusión")
add_body(doc,
    "Los ocho hitos conforman una secuencia completa de implementación. Las primeras seis semanas establecieron la "
    "base de operación, los accesos y las reglas comunes. Las diez semanas restantes incorporaron los procesos de "
    "prospectos, asociados, membresías, pagos, cobranza, documentos y reportes.")
add_body(doc,
    "Al cierre de las 16 semanas referenciales, la organización cuenta con una plataforma que integra información, "
    "reduce tareas dispersas y mejora la visibilidad sobre la relación con prospectos y asociados. El producto base "
    "previsto en los ocho hitos principales quedó cubierto, sin considerar las actividades de subsanación.")

doc.add_page_break()
add_level1(doc, "Referencias")
reference_titles = [
    "Hito 1: Base técnica del proyecto",
    "Hito 2: Usuarios, roles y configuración base",
    "Hito 3: Catálogos y reglas maestras",
    "Hito 4: Módulo de prospectos",
    "Hito 5: Conversión a asociado y ficha principal",
    "Hito 6: Membresías, pagos y cobranza",
    "Hito 7: Gestión documental y almacenamiento",
    "Hito 8: Reportes, exportaciones y automatizaciones",
]
for idx, ref_title in enumerate(reference_titles):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    font_run(p.add_run(f"Equipo del Sistema de Asociados. (s. f.-{suffixes[idx]}). "))
    font_run(p.add_run(ref_title), italic=True)
    font_run(p.add_run(" [Documento interno]."))

doc.core_properties.title = "Hitos realizados del Sistema de Asociados"
doc.core_properties.subject = "Informe ejecutivo en formato APA 7"
doc.core_properties.author = "Equipo del Sistema de Asociados"
doc.core_properties.keywords = "APA 7, hitos, sistema de asociados, informe ejecutivo"
doc.save(OUT)
print(OUT.resolve())
